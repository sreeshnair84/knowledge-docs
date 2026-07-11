"""Convert DOCX files to Markdown, preserving headings, lists, tables, and bold/italic."""

import io
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def table_to_md(table) -> str:
    rows = [[cell_text(c) for c in row.cells] for row in table.rows]
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * col_count) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return "\n".join(filter(None, [header, sep, body]))


def run_text(run) -> str:
    text = run.text
    if not text.strip():
        return text
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def para_to_md(para) -> str:
    style = para.style.name if para.style else ""
    inline = "".join(run_text(r) for r in para.runs)

    if style.startswith("Heading 1"):
        return f"# {inline.strip()}"
    if style.startswith("Heading 2"):
        return f"## {inline.strip()}"
    if style.startswith("Heading 3"):
        return f"### {inline.strip()}"
    if style.startswith("Heading 4"):
        return f"#### {inline.strip()}"
    if style.startswith("Heading 5"):
        return f"##### {inline.strip()}"

    # Bullet / list
    num_pr = para._element.find(qn("w:numPr"))
    if num_pr is not None:
        ilvl_el = num_pr.find(qn("w:ilvl"))
        level = int(ilvl_el.get(qn("w:val"), 0)) if ilvl_el is not None else 0
        indent = "  " * level
        return f"{indent}- {inline.strip()}"

    return inline


def open_doc(docx_path: Path) -> Document:
    """Handle files that have Markdown frontmatter prepended to DOCX binary."""
    raw = docx_path.read_bytes()
    # ZIP magic bytes for DOCX
    pk_offset = raw.find(b"PK\x03\x04")
    if pk_offset == -1:
        pk_offset = raw.find(b"PK\x05\x06")
    if pk_offset > 0:
        raw = raw[pk_offset:]
    return Document(io.BytesIO(raw))


def extract_frontmatter(docx_path: Path) -> str:
    """Pull the YAML frontmatter block prepended before the DOCX binary, if present."""
    try:
        raw = docx_path.read_bytes()
        if not raw.startswith(b"---"):
            return ""
        text = raw.split(b"PK\x03\x04")[0].decode("utf-8", errors="replace")
        return text.strip()
    except Exception:
        return ""


def convert(docx_path: Path, out_path: Path):
    doc = open_doc(docx_path)
    frontmatter = extract_frontmatter(docx_path)
    if not frontmatter:
        slug = docx_path.stem.replace("_", " ").replace("-", " ").title()
        frontmatter = f'---\ntitle: "{slug}"\n---'

    lines = [frontmatter, ""]

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "tbl":
            # Find the table object matching this element
            for tbl in doc.tables:
                if tbl._element is child:
                    lines.append(table_to_md(tbl))
                    lines.append("")
                    break
        elif tag == "p":
            # Find matching paragraph
            for para in doc.paragraphs:
                if para._element is child:
                    md = para_to_md(para)
                    lines.append(md)
                    break

    # Clean up excessive blank lines
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)

    out_path.write_text(text.strip() + "\n", encoding="utf-8")
    print(f"  OK  {docx_path.name} -> {out_path.name}")


DOCX_FILES = [
    "docs/ai-development/aidlc/Agile_in_the_Age_of_Agentic_AI_2026.docx",
    "docs/ai-development/testing/Agent_Testing_Monitoring_Evaluation.docx",
    "docs/ai-usecases/eu-bank-ai-copilot-architecture.docx",
    "docs/knowledge-engineering/knowledge/Complex_RAG_Deep_Dive.docx",
    "docs/coding-tools/claude/Anthropic_Partner_Cert_Study_Guide.docx",
    "docs/coding-tools/claude/CCAF_Study_Guide.docx",
    "docs/coding-tools/claude/CCAF_Advanced_Supplement.docx",
    "docs/coding-tools/claude/claude-best-practices.docx",
    "docs/enterprise-architecture/specialization/AI_Invoice_Auditor_Architecture_v5.docx",
    "docs/soft-skills/CTO_Voice_Mastery_Program.docx",
    "docs/soft-skills/Mental_Models_for_Voice_Training.docx",
    "docs/soft-skills/Storytelling_Exercise_Workbook.docx",
    "docs/soft-skills/Voice_Training_30Day_Plan.docx",
    "docs/agentic-systems/platform/agentic_platform_bestpractices.docx",
    "docs/interview-prep/AI_Engineer_Question_Bank.docx",
    "docs/interview-prep/Hard_Scenarios_Interview_Prep.docx",
    "docs/interview-prep/ML_AI_Interview_Mastery_Guide.docx",
]


def main():
    root = Path(__file__).parent.parent
    for rel in DOCX_FILES:
        docx_path = root / rel
        if not docx_path.exists():
            print(f" MISS {rel}")
            continue
        out_path = docx_path.with_suffix(".md")
        if out_path.exists():
            print(f" SKIP {out_path.name} (already exists)")
            continue
        try:
            convert(docx_path, out_path)
        except Exception as e:
            print(f"  ERR {docx_path.name}: {e}")


if __name__ == "__main__":
    main()
